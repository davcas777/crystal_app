"""Extract usable content from uploaded files.

Returns a dict shaped like:
    {"kind": "text",  "text": "..."}   for documents / spreadsheets
    {"kind": "image", "mime": "...", "b64": "..."}   for images

Images are base64-encoded so multimodal endpoints (Claude, GPT-4o) can ingest
them via the OpenAI vision payload schema.  Text from PDFs / Word / CSV is
inlined into the prompt.
"""

from __future__ import annotations

import base64
import io
import mimetypes
from typing import Any


def _safe_truncate(text: str, max_chars: int = 60_000) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n[... archivo truncado a {max_chars} caracteres ...]"


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            txt = ""
        if txt.strip():
            pages.append(f"[Página {i}]\n{txt.strip()}")
    return "\n\n".join(pages) if pages else "(PDF sin texto extraíble — puede ser escaneado.)"


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts) if parts else "(Documento sin contenido extraíble.)"


def extract_file_content(uploaded_file: Any) -> dict:
    """Inspect an UploadedFile and return its model-ready payload."""
    name = uploaded_file.name
    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    data = uploaded_file.getvalue()
    mime = uploaded_file.type or mimetypes.guess_type(name)[0] or "application/octet-stream"

    image_suffixes = {"png", "jpg", "jpeg", "webp", "gif"}
    if suffix in image_suffixes or mime.startswith("image/"):
        return {
            "kind": "image",
            "mime": mime if mime.startswith("image/") else f"image/{suffix}",
            "b64": base64.b64encode(data).decode("ascii"),
            "name": name,
        }

    try:
        if suffix == "pdf":
            text = _extract_pdf(data)
        elif suffix in {"docx", "doc"}:
            text = _extract_docx(data)
        elif suffix in {"txt", "md", "csv"}:
            text = data.decode("utf-8", errors="replace")
        else:
            text = data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        text = f"(No se pudo leer el archivo {name}: {exc})"

    return {"kind": "text", "text": _safe_truncate(text), "name": name}
